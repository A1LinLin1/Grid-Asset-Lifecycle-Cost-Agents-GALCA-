import os
import random
import pandas as pd
from datetime import datetime, timedelta
from docx import Document

# 设置基础目录
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DATA_DIR = os.path.join(os.path.dirname(BASE_DIR), "data")
os.makedirs(DATA_DIR, exist_ok=True)

# 1. 业务字典配置
EQUIPMENT_CONFIG = {
    "电缆终端": {"models": ["YJZWC4-110", "YJZWI4-220"], "base_cost": 15.0, "vendor": "XX线缆集团"},
    "交流避雷器": {"models": ["Y5WZ-17/45", "Y10WZ-51"], "base_cost": 2.5, "vendor": "XX防雷设备厂"},
    "高压熔断器": {"models": ["XRNP-12", "RN2-35"], "base_cost": 0.5, "vendor": "XX电器制造局"},
    "隔离开关": {"models": ["GW4-126", "GN30-12"], "base_cost": 8.0, "vendor": "XX高压开关厂"}
}

DEPARTMENTS = ["运维中心", "输变电检修部", "自动化运行部", "物资部"]

def random_date(start_year=2020, end_year=2025, seasonal=False):
    """生成随机日期，支持雷雨季节性(针对避雷器)"""
    start = datetime(start_year, 1, 1)
    end = datetime(end_year, 12, 31)
    
    if seasonal and random.random() < 0.7:
        # 70% 的概率发生在 6-9 月（雷雨季）
        year = random.randint(start_year, end_year)
        month = random.randint(6, 9)
        day = random.randint(1, 28)
        return datetime(year, month, day)
        
    return start + timedelta(days=random.randint(0, (end - start).days))

def generate_work_orders():
    """生成基于 LCC 浴盆曲线物理规律的工单数据"""
    records = []
    wo_counter = 1000
    
    for eq_type, config in EQUIPMENT_CONFIG.items():
        # 每种设备生成 150-200 条工单
        num_orders = random.randint(150, 200)
        for _ in range(num_orders):
            date_obj = random_date(seasonal=(eq_type == "交流避雷器"))
            year = date_obj.year
            model = random.choice(config["models"])
            
            # --- 核心逻辑：注入 LCC 规律 ---
            # 基础概率
            prob = random.random()
            task_type = "运行" # 默认
            cost = random.uniform(0.1, 1.0) # 默认极低成本
            summary = "日常巡视"
            
            # 随着时间推移（老化），设备发生故障和大修的概率增加
            aging_factor = (year - 2020) / 5.0 # 0.0 到 1.0
            
            if eq_type == "高压熔断器":
                # 熔断器绝大多数是直接熔断更换（故障成本）
                if prob < 0.8:
                    task_type, summary = "故障处理", "熔断器熔断更换"
                    cost = config["base_cost"] * random.uniform(0.8, 1.2)
            elif eq_type == "电缆终端":
                # 电缆终端后期维护和故障成本极高
                if prob < 0.4 + aging_factor * 0.4:
                    task_type, summary = "检修", "绝缘介损测试与修复"
                    cost = random.uniform(3.0, 8.0) * (1 + aging_factor) # 成本随老化逐年上升
                elif prob < 0.5 + aging_factor * 0.4:
                    task_type, summary = "故障处理", "终端击穿抢修"
                    cost = random.uniform(10.0, 25.0)
            elif eq_type == "隔离开关":
                # 隔离开关后期面临机械卡涩大修
                if year >= 2023 and prob < 0.6:
                    task_type, summary = "维保", "传动机构大修除锈"
                    cost = random.uniform(2.0, 5.0)
            elif eq_type == "交流避雷器":
                # 避雷器主要是雷击损坏
                if 6 <= date_obj.month <= 9 and prob < 0.7:
                    task_type, summary = "故障处理", "雷击击穿更换"
                    cost = config["base_cost"] * random.uniform(0.9, 1.3)
                    
            records.append({
                "公司": "A电力公司",
                "设备类型": eq_type,
                "工单编号": f"WO{date_obj.strftime('%Y%m')}-{wo_counter}",
                "日期": date_obj.strftime("%Y-%m-%d"),
                "型号": model,
                "任务类型": task_type,
                "内容摘要": f"{eq_type} - {summary}",
                "工期(天)": random.randint(1, 5),
                "成本(万元)": round(cost, 2),
                "执行部门": random.choice(DEPARTMENTS),
                "状态": random.choice(["已完成", "已验收", "执行中"])
            })
            wo_counter += 1
            
    # 转为 DataFrame 并按时间排序
    df = pd.DataFrame(records)
    df = df.sort_values(by="日期")
    excel_path = os.path.join(DATA_DIR, "LCC工单记录表_2020_2025.xlsx")
    df.to_excel(excel_path, index=False)
    print(f"✅ 成功生成符合物理规律的工单台账: {excel_path} (共 {len(df)} 条)")

def generate_contracts():
    """生成初始投资的 Word 采购合同"""
    for eq_type, config in EQUIPMENT_CONFIG.items():
        # 每年生成一份采购合同
        for year in range(2020, 2026):
            doc = Document()
            doc.add_heading(f'{year}年度 {eq_type} 集中采购协议', 0)
            
            contract_id = f"CG-{year}-{eq_type[:2]}-{random.randint(100, 999)}"
            # 采购数量和金额带有随机波动
            qty = random.randint(20, 100)
            total_amt = qty * config["base_cost"] * 10000 # 转为元
            
            doc.add_paragraph(f'项目名称：{year}年度配网物资统一采购批次')
            doc.add_paragraph(f'合同编号：{contract_id}')
            doc.add_paragraph(f'甲方：A电力公司物资部')
            doc.add_paragraph(f'乙方：{config["vendor"]}')
            
            doc.add_heading('一、 核心条款', level=1)
            doc.add_paragraph(f'1. 采购设备：{eq_type}，规格型号涵盖 {", ".join(config["models"])}。')
            doc.add_paragraph(f'2. 采购数量：{qty} 批/套。')
            doc.add_paragraph(f'3. 合同总金额：人民币 {total_amt:,.2f} 元。')
            doc.add_paragraph(f'4. 签署日期：{year}-01-15。')
            
            file_path = os.path.join(DATA_DIR, f"{year}年_{eq_type}_采购合同.docx")
            doc.save(file_path)
    print(f"✅ 成功生成四类设备 2020-2025 年的 Word 采购合同。")

if __name__ == "__main__":
    generate_work_orders()
    generate_contracts()