import pandas as pd
from unstructured.partition.auto import partition
from langchain_core.documents import Document

class UniversalLoader:
    """万能文档加载器：支持 PDF, DOCX, XLSX, JPG/PNG"""
    
    def __init__(self, file_path: str):
        self.file_path = file_path

    def load_as_dataframe(self):
        """专门针对 Excel 和 CSV 的表格数据"""
        if self.file_path.endswith(('.xlsx', '.xls')):
            return pd.read_excel(self.file_path)
        elif self.file_path.endswith('.csv'):
            return pd.read_csv(self.file_path)
        return None

    def load_as_text(self):
        """针对 PDF, Word 和 图片的非结构化提取"""
        # 1. 处理 Word 文档 (.docx)
        if self.file_path.endswith('.docx'):
            import docx
            doc = docx.Document(self.file_path)
            content = "\n".join([para.text for para in doc.paragraphs if para.text.strip() != ""])
            return Document(page_content=content, metadata={"source": self.file_path})
            
        # 2. 处理纯文本 (.txt)
        elif self.file_path.endswith('.txt'):
            with open(self.file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            return Document(page_content=content, metadata={"source": self.file_path})
            
        # 3. 这里可以保留你之前的 unstructured 逻辑处理 PDF 等其他格式
        else:
            from unstructured.partition.auto import partition
            elements = partition(filename=self.file_path)
            content = "\n\n".join([str(el) for el in elements])
            return Document(page_content=content, metadata={"source": self.file_path})