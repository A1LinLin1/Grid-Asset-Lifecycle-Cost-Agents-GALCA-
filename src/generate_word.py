import docx
import os

os.makedirs('data', exist_ok=True)

def create_word_contract():
    doc = docx.Document()
    doc.add_heading('高压电缆采购与施工合同', 0)
    
    doc.add_paragraph('项目名称：2025年城网改造高压电缆入地工程')
    doc.add_paragraph('合同编号：WORD-2025-001')
    doc.add_paragraph('甲方：国家电网XX供电局')
    doc.add_paragraph('乙方：XX线缆制造与施工总公司')
    
    doc.add_heading('一、 核心条款', level=1)
    doc.add_paragraph('1. 标的物：110kV 交联聚乙烯绝缘电力电缆及配套附件。')
    doc.add_paragraph('2. 合同总金额：人民币 3,500,000.00 元（大写：叁佰伍拾万元整）。')
    doc.add_paragraph('3. 签署日期：2025-01-20。')
    
    doc.save('data/sample_contract.docx')
    print("✅ 已生成模拟 Word 合同: data/sample_contract.docx")

if __name__ == "__main__":
    create_word_contract()